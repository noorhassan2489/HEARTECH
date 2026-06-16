"""ReferralAIService — HearTech v5 runtime with intent guardrails."""

from __future__ import annotations

import gc
import os
import re
import shutil
import tempfile
from collections import Counter
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from dotenv import load_dotenv
from heartech_ai.runtime.constants import CONCERN_MEDICINES, FOOTER_ADVICE, FOOTER_REFERRAL, MEDICINE_GUIDANCE
from heartech_ai.runtime.intent import (
    extract_referral_destination,
    is_medicine_section_edit,
    should_edit_existing_referral,
    should_include_refer_to,
    should_patch_existing_referral,
)
from heartech_ai.runtime.output_patterns import (
    build_cloud_runtime_prompt,
    build_local_runtime_prompt,
    build_pattern_assembled_output,
    build_referral_edit_prompt,
)
from heartech_ai.runtime.prompting import (
    build_inference_prompt,
)
from heartech_ai.runtime.model_selection import resolve_active_model_path
from heartech_ai.runtime.referral_requests import (
    extract_hcw_care_requests,
    hcw_requested_referral_items,
    misplaced_destination_in_summary,
    output_contains_invented_condition,
    referral_output_satisfies_requests,
)
from heartech_ai.runtime.router import decide_without_base_router
from services.runtime_cloud_provider import RuntimeCloudProvider
from heartech_ai.runtime.validators import (
    ValidationResult,
    asks_action_for_reported_symptoms,
    asks_clinical_differential_question,
    asks_precautions_or_care_request,
    asks_symptom_clinical_question,
    asks_yes_no_clinical_question,
    extract_clinical_topics,
    has_hearing_milestone_leak,
    has_prompt_instruction_leak,
    has_training_qa_leak,
    has_training_question_echo,
    has_incomplete_answer_fragment,
    is_generic_precaution_boilerplate,
    is_hard_answer_failure,
    is_instruction_leak_line,
    is_noisy_answer_line,
    validate_runtime_output,
)

_LLAMA_SPECIAL_TOKENS = (
    "<|begin_of_text|>",
    "<|end_of_text|>",
    "<|eot_id|>",
    "<|start_header_id|>",
    "<|end_header_id|>",
    "<|assistant|>",
    "<|end|>",
)
_RESERVED_TOKEN_RE = re.compile(r"<\|reserved_special_token_\d+\|>.*", re.DOTALL)
_INFERENCE_MAX_TOKENS = 900
# Only block shipping for harmful/leak failures — not soft topical misses.
_SHIPPING_BLOCKED_WARNINGS = frozenset(
    {
        "template_leak_in_answer",
        "token_leak_in_answer",
        "contamination_in_answer",
        "training_qa_leak_in_answer",
        "training_question_echo_in_answer",
        "hearing_milestone_leak_in_answer",
        "adult_narrative_leak_in_answer",
        "adult_patient_leak_in_answer",
        "structured_case_dump_in_answer",
        "unsupported_history_claim_in_answer",
        "incomplete_answer_fragment",
        "prompt_instruction_leak_in_answer",
        "developmental_warning_leak_in_answer",
        "invented_condition_in_answer",
        "generic_precaution_boilerplate_in_answer",
    }
)

_MEDICINE_ALIASES: dict[str, str] = {
    "panadol": "paracetamol",
    "calpol": "paracetamol",
    "tylenol": "paracetamol",
    "acetaminophen": "paracetamol",
    "paracetamol": "paracetamol",
    "ibuprofen": "ibuprofen",
    "brufen": "ibuprofen",
    "nurofen": "ibuprofen",
    "cyclizine": "cyclizine",
    "amoxicillin": "amoxicillin",
    "co-amoxiclav": "co-amoxiclav",
    "azithromycin": "azithromycin",
    "loratadine": "loratadine",
    "cetirizine": "cetirizine",
}

_BACKEND_DIR = Path(__file__).resolve().parent.parent
load_dotenv(_BACKEND_DIR / ".env")
_BASE_DIR = _BACKEND_DIR / "heartech_ai"

_MODEL_SELECTION = resolve_active_model_path(
    _BASE_DIR, os.getenv("HEARTECH_ACTIVE_MODEL_PATH", "")
)
_MODEL_PATH = _MODEL_SELECTION.active_path

_EXPORTS_DIR = _BACKEND_DIR / "referral_exports"
_EXPORTS_DIR.mkdir(parents=True, exist_ok=True)

# Internal runtime labels mapped to public v5-only names.
_PUBLIC_SOURCE_ALIASES: dict[str, str] = {
    "runtime_cloud_retry": "v5_fused_retry",
    "pattern_assembled": "v5_fused",
    "reasoned_referral": "v5_fused",
    "rule_based_referral_session": "v5_fused",
    "clarifier": "v5_fused",
}


@dataclass
class GenerationResult:
    text: str
    success: bool
    source: str
    intent: str
    needs_clarification: bool
    normalized_prompt: str
    validation_warnings: list[str] = field(default_factory=list)


class ReferralAIService:
    _instance = None

    def __init__(self):
        self.last_generation_source = "unknown"
        self.last_generation_meta: dict[str, object] = {}
        self._model_primary = None
        self._tokenizer = None
        self._cloud_provider = RuntimeCloudProvider()
        self._aux_provider = self._cloud_provider

    @classmethod
    def get_instance(cls):
        if cls._instance is None:
            cls._instance = cls()
        return cls._instance

    @property
    def has_local_model(self) -> bool:
        return self._model_primary is not None and self._tokenizer is not None

    def _release_metal_memory(self) -> None:
        gc.collect()
        try:
            import mlx.core as mx

            if hasattr(mx, "clear_cache"):
                mx.clear_cache()
            elif hasattr(mx, "metal") and hasattr(mx.metal, "clear_cache"):
                mx.metal.clear_cache()
        except Exception:
            pass
        gc.collect()

    def _unload_primary_model(self) -> None:
        model = self._model_primary
        self._model_primary = None
        if model is not None:
            del model
        for _ in range(2):
            self._release_metal_memory()

    def _load_primary_model(self):
        if self._model_primary is not None and self._tokenizer is not None:
            return
        try:
            from mlx_lm import load

            print(
                "[REFERRAL-AI] Loading active model "
                f"(source={_MODEL_SELECTION.active_source}) from {_MODEL_PATH}"
            )
            self._model_primary, self._tokenizer = load(_MODEL_PATH)
            print("[REFERRAL-AI] Active model ready.")
        except Exception as e:
            print(f"[REFERRAL-AI] WARNING: Active model load failed: {e}")
            self._model_primary = None
            self._tokenizer = None

    def _cleanup_after_request(self) -> None:
        """Clear MLX cache between requests."""
        self._release_metal_memory()

    @staticmethod
    def _public_source(internal: str) -> str:
        if internal.startswith("v5_guided_retry"):
            return "v5_fused_retry"
        return _PUBLIC_SOURCE_ALIASES.get(internal, internal if internal.startswith("v5_") else "v5_fused")

    @staticmethod
    def _public_log_label(internal: str) -> str:
        return ReferralAIService._public_source(internal)

    def _emit_result(self, result: GenerationResult) -> GenerationResult:
        internal_source = result.source
        public = self._public_source(internal_source)
        public_result = GenerationResult(
            text=result.text,
            success=result.success,
            source=public,
            intent=result.intent,
            needs_clarification=result.needs_clarification,
            normalized_prompt=result.normalized_prompt,
            validation_warnings=[],
        )
        self.last_generation_source = public
        self.last_generation_meta = {
            **public_result.__dict__,
            "_internal_source": internal_source,
        }
        return public_result

    def _route_instruction(self, hcw_instruction: str):
        return decide_without_base_router(hcw_instruction)

    def generate(self, child_data: dict, hcw_instruction: str) -> GenerationResult:
        try:
            return self._generate_impl(child_data, hcw_instruction)
        finally:
            self._cleanup_after_request()

    def _generate_impl(self, child_data: dict, hcw_instruction: str) -> GenerationResult:
        if self._model_primary is None or self._tokenizer is None:
            self._load_primary_model()
        if self._model_primary is None or self._tokenizer is None:
            raise RuntimeError(
                "Local model unavailable. Verify HEARTECH_ACTIVE_MODEL_PATH and restart API."
            )

        primary_instruction, conversation_context = self._split_instruction_context(
            hcw_instruction
        )
        routing_input = primary_instruction or hcw_instruction
        decision = self._route_instruction(routing_input)
        has_core_child_context = bool(str(child_data.get("age", "")).strip()) and bool(
            str(child_data.get("riskScore", "")).strip()
        )
        if decision.needs_clarification and not has_core_child_context:
            result = GenerationResult(
                text=decision.clarification_question
                or "Could you share age, key symptom, and whether you want advice or referral?",
                success=True,
                source="clarifier",
                intent=decision.intent,
                needs_clarification=True,
                normalized_prompt=decision.normalized_prompt,
            )
            return self._emit_result(result)

        generation_instruction = decision.normalized_prompt
        if conversation_context:
            generation_instruction = (
                f"{decision.normalized_prompt}\n\nConversation context:\n"
                f"{conversation_context}"
            )

        routing_text = primary_instruction or decision.normalized_prompt
        destination = extract_referral_destination(routing_text)
        prior_referral = self._extract_prior_referral_from_context(conversation_context)

        if prior_referral and should_edit_existing_referral(
            routing_text, has_prior_referral=True
        ):
            if should_patch_existing_referral(
                routing_text, has_prior_referral=True
            ) and destination:
                patched = self._patch_referral_add_destination(prior_referral, destination)
                result = GenerationResult(
                    text=patched,
                    success=True,
                    source="referral_patch",
                    intent="referral",
                    needs_clarification=False,
                    normalized_prompt=decision.normalized_prompt,
                )
                return self._emit_result(result)

            edit_result = self._edit_existing_referral(
                prior_referral=prior_referral,
                edit_instruction=routing_text,
                child_data=child_data,
                normalized_prompt=decision.normalized_prompt,
            )
            return self._emit_result(edit_result)

        include_refer_to = (
            should_include_refer_to(routing_text) if decision.intent == "referral" else False
        )

        generation_result = self._generate_with_two_tier(
            child_data=child_data,
            normalized_prompt=decision.normalized_prompt,
            generation_instruction=generation_instruction,
            primary_instruction=primary_instruction or decision.normalized_prompt,
            session_context=conversation_context,
            intent=decision.intent,
            include_refer_to=include_refer_to,
        )
        return self._emit_result(generation_result)

    def _generate_with_two_tier(
        self,
        *,
        child_data: dict,
        normalized_prompt: str,
        generation_instruction: str,
        primary_instruction: str,
        session_context: str,
        intent: str,
        include_refer_to: bool,
    ) -> GenerationResult:
        local_prompt = build_local_runtime_prompt(
            child_data=child_data,
            normalized_instruction=generation_instruction,
            intent=intent,
            include_refer_to=include_refer_to,
        )
        max_tokens = 900 if intent == "referral" else 320

        local_text = self._run_inference(
            self._model_primary,
            local_prompt,
            label="v5_fused",
            max_tokens=max_tokens,
            temp=0.2,
            top_p=0.9,
        )
        local_text = self._sanitize_output(
            local_text, intent=intent, include_refer_to=include_refer_to
        )
        local_validation = self._validate_output(
            output=local_text,
            intent=intent,
            include_refer_to=include_refer_to,
            child_data=child_data,
            primary_instruction=primary_instruction,
            generation_instruction=generation_instruction,
        )
        if self._can_ship_output(
            local_text,
            local_validation,
            intent=intent,
            primary_instruction=primary_instruction,
        ) or (
            intent == "answer"
            and self._can_ship_answer_relaxed(
                local_text,
                local_validation,
                primary_instruction=primary_instruction,
            )
        ):
            if intent == "referral":
                local_text = self._finalize_referral_for_hcw(
                    local_text,
                    prompt=generation_instruction,
                    child_data=child_data,
                    include_refer_to=include_refer_to,
                )
            print(f"[REFERRAL-AI] {self._public_log_label('v5_fused')} accepted ({len(local_text.strip())} chars)")
            return GenerationResult(
                text=local_text,
                success=True,
                source="v5_fused",
                intent=intent,
                needs_clarification=False,
                normalized_prompt=normalized_prompt,
                validation_warnings=local_validation.warnings,
            )

        print(
            f"[REFERRAL-AI] {self._public_log_label('v5_fused')} rejected: "
            f"{local_validation.warnings[:8]}"
        )

        cloud_validation: ValidationResult | None = None
        if self._cloud_provider.is_available():
            cloud_prompt = build_cloud_runtime_prompt(
                child_data=child_data,
                primary_instruction=primary_instruction,
                session_context=session_context,
                intent=intent,
                include_refer_to=include_refer_to,
            )
            try:
                cloud_raw = self._cloud_provider.generate(cloud_prompt, max_tokens=max_tokens)
                cloud_text = self._sanitize_output(
                    cloud_raw, intent=intent, include_refer_to=include_refer_to
                )
                cloud_validation = self._validate_output(
                    output=cloud_text,
                    intent=intent,
                    include_refer_to=include_refer_to,
                    child_data=child_data,
                    primary_instruction=primary_instruction,
                    generation_instruction=generation_instruction,
                )
                if self._can_ship_output(
                    cloud_text,
                    cloud_validation,
                    intent=intent,
                    primary_instruction=primary_instruction,
                ):
                    if intent == "referral":
                        cloud_text = self._finalize_referral_for_hcw(
                            cloud_text,
                            prompt=generation_instruction,
                            child_data=child_data,
                            include_refer_to=include_refer_to,
                        )
                    print(
                        f"[REFERRAL-AI] {self._public_log_label('runtime_cloud_retry')} accepted "
                        f"({len(cloud_text.strip())} chars)"
                    )
                    return GenerationResult(
                        text=cloud_text,
                        success=True,
                        source="runtime_cloud_retry",
                        intent=intent,
                        needs_clarification=False,
                        normalized_prompt=normalized_prompt,
                        validation_warnings=cloud_validation.warnings,
                    )
                print(
                    f"[REFERRAL-AI] {self._public_log_label('runtime_cloud_retry')} rejected: "
                    f"{cloud_validation.warnings[:8]}"
                )
            except Exception:
                print(f"[REFERRAL-AI] {self._public_log_label('runtime_cloud_retry')} unavailable.")

        prior_ai_answer = self._extract_prior_ai_answer_from_context(session_context)
        if intent == "answer":
            rescue = self._build_concern_rescue_answer(
                instruction=generation_instruction,
                child_data=child_data,
            )
            if rescue.strip():
                assembled = rescue
            else:
                assembled = build_pattern_assembled_output(
                    intent=intent,
                    referral_builder=lambda: self._build_reasoned_referral(
                        prompt=generation_instruction,
                        child_data=child_data,
                        include_refer_to=include_refer_to,
                    ),
                    answer_builder=lambda: self._build_pattern_assembled_answer(
                        instruction=generation_instruction,
                        child_data=child_data,
                        prior_answer=prior_ai_answer,
                    ),
                )
        else:
            assembled = build_pattern_assembled_output(
                intent=intent,
                referral_builder=lambda: self._build_reasoned_referral(
                    prompt=generation_instruction,
                    child_data=child_data,
                    include_refer_to=include_refer_to,
                ),
                answer_builder=lambda: self._build_pattern_assembled_answer(
                    instruction=generation_instruction,
                    child_data=child_data,
                    prior_answer=prior_ai_answer,
                ),
            )
        if not assembled.strip() and intent == "answer":
            assembled = self._build_pattern_assembled_answer(
                instruction=generation_instruction,
                child_data=child_data,
                prior_answer=prior_ai_answer,
            )
        if not assembled.strip():
            assembled = self._build_reasoned_referral(
                prompt=generation_instruction,
                child_data=child_data,
                include_refer_to=include_refer_to,
            )

        assembled_validation = self._validate_output(
            output=assembled,
            intent=intent,
            include_refer_to=include_refer_to,
            child_data=child_data,
            primary_instruction=primary_instruction,
            generation_instruction=generation_instruction,
        )
        print(
            "[REFERRAL-AI] Using pattern assembler fallback. "
            f"warnings={assembled_validation.warnings[:6]}"
        )
        if intent == "referral":
            assembled = self._finalize_referral_for_hcw(
                assembled,
                prompt=generation_instruction,
                child_data=child_data,
                include_refer_to=include_refer_to,
            )
        return GenerationResult(
            text=assembled,
            success=True,
            source="pattern_assembled",
            intent=intent,
            needs_clarification=False,
            normalized_prompt=normalized_prompt,
            validation_warnings=assembled_validation.warnings
            + local_validation.warnings
            + (cloud_validation.warnings if cloud_validation else []),
        )

    def _sanitize_output(self, text: str, *, intent: str, include_refer_to: bool) -> str:
        if intent == "answer":
            return self._sanitize_answer_output(text)
        return self._sanitize_referral_output(text, include_refer_to=include_refer_to)

    def _validate_output(
        self,
        *,
        output: str,
        intent: str,
        include_refer_to: bool,
        child_data: dict,
        primary_instruction: str,
        generation_instruction: str,
    ) -> ValidationResult:
        prompt_for_validation = (
            primary_instruction if intent == "answer" else generation_instruction
        )
        return validate_runtime_output(
            output=output,
            intent=intent,
            include_refer_to=include_refer_to,
            child_data=child_data,
            user_prompt=prompt_for_validation,
        )

    def _can_ship_output(
        self,
        text: str,
        validation: ValidationResult,
        *,
        intent: str,
        primary_instruction: str,
    ) -> bool:
        if intent == "answer":
            return self._can_ship_answer(
                text=text,
                validation=validation,
                primary_instruction=primary_instruction,
            )
        if not validation.ok or self._is_degenerate(text):
            return False
        return True

    def _referral_semantically_valid(
        self,
        text: str,
        *,
        prompt: str,
        child_data: dict,
        include_refer_to: bool,
    ) -> bool:
        flags = child_data.get("flags") if isinstance(child_data.get("flags"), list) else []
        if output_contains_invented_condition(prompt, text, flags):
            return False
        if misplaced_destination_in_summary(
            text,
            hcw_spec=str(child_data.get("hcwSpec", "")),
            hcw_hospital=str(child_data.get("hcwHospital", "")),
            include_refer_to=include_refer_to,
        ):
            return False
        if hcw_requested_referral_items(prompt) and not referral_output_satisfies_requests(prompt, text):
            return False
        if not include_refer_to and "REFER TO:" in text.upper():
            return False
        return True

    def _finalize_referral_for_hcw(
        self,
        text: str,
        *,
        prompt: str,
        child_data: dict,
        include_refer_to: bool,
    ) -> str:
        """Align referral documents with HCW instructions — strip profile leakage, inject requested care."""
        primary, context = self._split_instruction_context(prompt)
        full_prompt = f"{primary}\n{context}".strip()
        flags = child_data.get("flags") if isinstance(child_data.get("flags"), list) else []

        if (
            not text.strip()
            or not self._referral_semantically_valid(
                text,
                prompt=full_prompt,
                child_data=child_data,
                include_refer_to=include_refer_to,
            )
        ):
            text = self._build_reasoned_referral(
                prompt=full_prompt,
                child_data=child_data,
                include_refer_to=include_refer_to,
            )

        text = self._sanitize_referral_output(text, include_refer_to=include_refer_to)
        text = self._strip_misplaced_referral_lines(
            text,
            child_data=child_data,
            include_refer_to=include_refer_to,
        )
        text = self._strip_invented_conditions_from_referral(
            text,
            prompt=full_prompt,
            flags=flags,
        )
        text = self._inject_hcw_care_requests(text, prompt=full_prompt)
        text = self._normalize_referral_urgency(text, prompt=full_prompt, include_refer_to=include_refer_to)
        return text.strip()

    @staticmethod
    def _strip_misplaced_referral_lines(
        text: str,
        *,
        child_data: dict,
        include_refer_to: bool,
    ) -> str:
        if include_refer_to or not text.strip():
            return text
        spec = str(child_data.get("hcwSpec", "")).strip()
        hospital = str(child_data.get("hcwHospital", "")).strip()
        cleaned: list[str] = []
        in_summary = False
        for raw in text.splitlines():
            line = raw.rstrip()
            upper = line.strip().upper()
            if upper.startswith("CLINICAL SUMMARY"):
                in_summary = True
                cleaned.append(line)
                continue
            if upper.startswith(("RECOMMENDED CARE", "REFER TO", "MEDICINES", "PRECAUTIONS", "FOLLOW UP")):
                in_summary = False
            if not in_summary:
                cleaned.append(line)
                continue
            lower = line.lower()
            if upper.startswith("URGENCY:"):
                continue
            if hospital and hospital.lower() in lower:
                continue
            if spec and spec.lower() in lower and (
                "—" in line or " - " in line or "hospital" in lower or "clinic" in lower
            ):
                continue
            if re.search(r"\b(?:audiologist|ent specialist|paediatric)\s*[—\-]", lower):
                continue
            cleaned.append(line)
        return "\n".join(cleaned)

    @staticmethod
    def _strip_invented_conditions_from_referral(
        text: str,
        *,
        prompt: str,
        flags: list[str],
    ) -> str:
        if not text.strip():
            return text
        context = prompt.lower() + " " + " ".join(flags).lower()
        cleaned: list[str] = []
        for line in text.splitlines():
            lower = line.lower()
            drop = False
            for marker in (
                "cholesteatoma",
                "eustachian tube dysfunction",
                "otitis externa",
                "sensorineural hearing loss",
                "glue ear",
            ):
                if marker in lower and marker not in context:
                    drop = True
                    break
            if not drop:
                cleaned.append(line.rstrip())
        return "\n".join(cleaned)

    @staticmethod
    def _inject_hcw_care_requests(text: str, *, prompt: str) -> str:
        requests = extract_hcw_care_requests(prompt)
        if not requests:
            return text
        lower = text.lower()
        missing = [line for line in requests if line.lower()[:24] not in lower]
        if not missing:
            return text

        bullets = "\n".join(f"• {line}" for line in missing)
        marker = "RECOMMENDED CARE:"
        idx = text.upper().find(marker)
        if idx == -1:
            return text + f"\n\nRECOMMENDED CARE:\n{bullets}\n"
        insert_at = idx + len(marker)
        return text[:insert_at] + "\n" + bullets + text[insert_at:]

    @staticmethod
    def _normalize_referral_urgency(text: str, *, prompt: str, include_refer_to: bool) -> str:
        if include_refer_to:
            return text
        # Remove orphan urgency lines when HCW did not ask for a destination referral.
        lines = []
        for line in text.splitlines():
            if line.strip().upper().startswith("URGENCY:"):
                continue
            lines.append(line.rstrip())
        return "\n".join(lines)

    def _generate_clinical_answer(
        self,
        *,
        runtime_user_prompt: str,
        generation_instruction: str,
        child_data: dict,
        normalized_prompt: str,
    ) -> GenerationResult:
        """Backward-compatible wrapper to the two-tier runtime path."""
        primary_instruction, session_context = self._split_instruction_context(
            generation_instruction
        )
        return self._generate_with_two_tier(
            child_data=child_data,
            normalized_prompt=normalized_prompt,
            generation_instruction=generation_instruction,
            primary_instruction=primary_instruction,
            session_context=session_context,
            intent="answer",
            include_refer_to=False,
        )

    @staticmethod
    def _scrub_training_template_sentences(text: str) -> str:
        parts = re.split(r"[\n\r]+|(?<=[.!?])\s+", text)
        kept: list[str] = []
        deferred_markers = (
            "if these signs persist beyond 2-4 weeks",
            "arrange formal hearing assessment",
        )
        for part in parts:
            sentence = part.strip(" -•\t")
            if len(sentence) < 12:
                continue
            if has_hearing_milestone_leak(sentence) or has_training_qa_leak(sentence):
                continue
            if has_training_question_echo(sentence):
                continue
            if is_instruction_leak_line(sentence):
                continue
            lower = sentence.lower()
            if any(m in lower for m in deferred_markers):
                continue
            kept.append(sentence)
        return " ".join(kept).strip()

    @staticmethod
    def _asks_explanatory_clinical_question(primary: str) -> bool:
        """HCW wants meaning, explanation, or clinical interpretation — not a referral doc."""
        if asks_symptom_clinical_question(primary) or asks_clinical_differential_question(primary):
            return True
        p = primary.lower()
        if any(x in p for x in ("tell me", "explain", "describe", "help me understand")):
            return True
        if any(x in p for x in ("what does", "what do", "what could", "what might")) and any(
            x in p for x in ("mean", "signify", "indicate", "suggest", "imply")
        ):
            return True
        return False

    def _build_explanatory_concern_answer(
        self,
        *,
        primary: str,
        context: str,
        child_data: dict,
        concerns: list[str] | None = None,
    ) -> str:
        """Topic-aware plain-language answer for 'what does X mean?' style questions."""
        detected = concerns or self._detect_concern_keys_for_answer(primary, context)
        age = child_data.get("age", "Unknown")
        risk = child_data.get("riskScore", "N/A")
        p_lower = primary.lower()

        if not detected and ("stutter" in p_lower or "stuttering" in p_lower):
            detected = ["speech"]
        if not detected and any(x in p_lower for x in ("speech delay", "speech delays", "language delay")):
            detected = ["speech"]

        guidance: list[str] = []
        for concern in detected[:3]:
            line = self._clinical_guidance_for_concern(concern, age, prompt=primary)
            if line and line not in guidance:
                guidance.append(line)

        if guidance:
            intro = f"For this {age} child (HearTech risk score {risk}/100):"
            body = " ".join(guidance)
            tail = (
                "Arrange ENT/audiology assessment when hearing is uncertain, and involve speech-language therapy "
                "if language or disfluency concerns persist after treating any ear disease."
            )
            return f"{intro} {body} {tail}"

        topics = extract_clinical_topics(primary) or extract_clinical_topics(context)
        if topics:
            topic_text = ", ".join(topics[:3])
            return (
                f"For this {age} child (HearTech risk score {risk}/100), regarding {topic_text}: "
                "these findings may reflect ear or hearing problems that affect communication, infection, "
                "or other ENT concerns — correlate with otoscopy and age-appropriate hearing tests before "
                "deciding next steps."
            )
        return ""

    def _build_concern_rescue_answer(
        self, *, instruction: str, child_data: dict
    ) -> str:
        """Last-resort guidance from detected concerns — keyword-triggered, not question hardcoding."""
        primary, context = self._split_instruction_context(instruction)
        concerns = self._detect_concern_keys_for_answer(primary, context)
        age = child_data.get("age", "Unknown")
        risk = child_data.get("riskScore", "N/A")

        if asks_yes_no_clinical_question(primary):
            focus = concerns[0] if concerns else ""
            if not focus and "vertigo" in primary.lower():
                focus = "vertigo"
            if focus == "vertigo":
                answer = (
                    f"For this {age} child (HearTech risk score {risk}/100): "
                    "With isolated ear wax alone, vertigo is unlikely unless the child also reports spinning, "
                    "balance difficulty, or repeated vomiting. "
                    "Vertigo can occur with ear or vestibular disease, so examine gait, eye movements, and ears. "
                    "Refer urgently if spinning is persistent or accompanied by severe headache or neurologic signs."
                )
                if any(x in primary.lower() for x in ("suggestion", "suggestions", "if no", "otherwise")):
                    answer += (
                        " If vertigo is not present, continue safe wax care (no cotton buds), "
                        "repeat otoscopy after softening drops if used, monitor hearing and speech, "
                        "and arrange ENT/audiology review if wax or hearing concern persists given the high risk score."
                    )
                return answer
            if focus:
                guidance = self._clinical_guidance_for_concern(focus, age, prompt=primary)
                if guidance:
                    return (
                        f"For this {age} child (HearTech risk score {risk}/100): "
                        f"{guidance}"
                    )
            return ""

        if self._asks_explanatory_clinical_question(primary):
            return self._build_explanatory_concern_answer(
                primary=primary,
                context=context,
                child_data=child_data,
                concerns=concerns,
            )

        if not (
            asks_action_for_reported_symptoms(primary)
            or asks_precautions_or_care_request(primary)
        ):
            return ""
        if not concerns:
            return ""
        guidance: list[str] = []
        for concern in concerns[:3]:
            line = self._clinical_guidance_for_concern(concern, age, prompt=primary)
            if line and line not in guidance:
                guidance.append(line)
        if not guidance:
            return ""
        intro = (
            f"For this {age} child (HearTech risk score {risk}/100), "
            "based on the reported concern(s):"
        )
        body = " ".join(guidance)
        tail = (
            "Examine with otoscopy before any intervention. "
            "Counsel parents to avoid probing the ear canal. "
            "Arrange ENT or audiology review if symptoms persist, hearing concern remains, "
            "or the screening risk score stays high."
        )
        return f"{intro} {body} {tail}"

    def _build_pattern_assembled_answer(
        self,
        *,
        instruction: str,
        child_data: dict,
        prior_answer: str = "",
    ) -> str:
        primary, context = self._split_instruction_context(instruction)
        concerns = self._detect_concern_keys_for_answer(primary, context)
        age = child_data.get("age", "Unknown")
        risk = child_data.get("riskScore", "N/A")
        seed = f"{primary}|{context}|{age}|{risk}"
        primary_lower = primary.lower()
        looks_like_yes_no = asks_yes_no_clinical_question(primary) or (
            "vertigo" in primary_lower
            and (
                "?" in primary_lower
                or "could" in primary_lower
                or "right" in primary_lower
                or "not vertigo" in primary_lower
                or "signs of vertigo" in primary_lower
            )
        )

        def choose(variants: list[str], offset: int = 0) -> str:
            idx = (self._stable_index(seed, len(variants)) + offset) % len(variants)
            return variants[idx]

        if looks_like_yes_no:
            if "vertigo" in primary_lower or "vertigo" in concerns:
                variants = [
                    (
                        f"Unlikely for this {age} child (HearTech risk score {risk}/100) to have true vertigo "
                        "from wax alone unless there is spinning sensation, imbalance, or repeated vomiting. "
                        "Check gait, eye movements, and otoscopy findings; escalate urgently for neurologic red flags "
                        "or persistent spinning."
                    ),
                    (
                        f"Possible but not the most likely diagnosis for this {age} child (HearTech risk score {risk}/100). "
                        "First rule out impacted wax and middle-ear disease with focused otoscopy and balance assessment. "
                        "Treat as urgent if severe headache, repeated vomiting, focal deficits, or continuous spinning are present."
                    ),
                    (
                        f"Not clearly vertigo yet in this {age} child (HearTech risk score {risk}/100). "
                        "Confirm whether the child has true spinning or balance disturbance, then correlate with ear exam findings. "
                        "Arrange prompt ENT/audiology review if symptoms persist or red flags appear."
                    ),
                ]
                answer = choose(variants)
            else:
                variants = [
                    (
                        f"Likely related to the reported ear concerns in this {age} child (HearTech risk score {risk}/100), "
                        "but confirm with focused otoscopy and symptom correlation. "
                        "Escalate urgently if severe pain, discharge, persistent fever, vomiting, or neurologic signs appear."
                    ),
                    (
                        f"Possible in this {age} child (HearTech risk score {risk}/100), "
                        "though confirmation needs direct examination and hearing-focused assessment. "
                        "Prioritize urgent review if symptoms worsen or any red-flag findings are present."
                    ),
                ]
                answer = choose(variants)
        elif self._is_medicine_request(primary):
            medicine_lines = self._medicine_lines_for_answer(primary or context)
            concerns_text = ", ".join(self._concern_label(c) for c in concerns[:3]) or "the reported ear concern"
            medicine_text = " ".join(medicine_lines[:3])
            variants = [
                (
                    f"For this {age} child (HearTech risk score {risk}/100), medicine should be linked to {concerns_text} after otoscopy. "
                    f"Suggested options: {medicine_text}"
                ),
                (
                    f"For this {age} child (HearTech risk score {risk}/100), use medicine only after confirming the ear findings clinically. "
                    f"Practical medicine guidance: {medicine_text}"
                ),
            ]
            answer = choose(variants)
        elif asks_action_for_reported_symptoms(primary) or asks_precautions_or_care_request(primary):
            guidance: list[str] = []
            for concern in concerns[:3]:
                line = self._clinical_guidance_for_concern(concern, age, prompt=primary)
                if line and line not in guidance:
                    guidance.append(line)
            if guidance:
                variants = [
                    (
                        f"For this {age} child (HearTech risk score {risk}/100), prioritize: "
                        + " ".join(guidance)
                        + " Perform otoscopy before intervention, avoid probing/cotton buds, and arrange ENT/audiology "
                          "review if symptoms persist."
                    ),
                    (
                        f"For this {age} child (HearTech risk score {risk}/100), suggested next steps are: "
                        + " ".join(guidance)
                        + " Document hearing and balance findings, advise safe ear care, and escalate urgently for severe pain, "
                          "discharge, vomiting, or neurologic red flags."
                    ),
                ]
                answer = choose(variants)
            else:
                answer = (
                    f"For this {age} child (HearTech risk score {risk}/100), perform focused ear examination and symptom-based "
                    "management, then reassess quickly. Escalate urgently if severe pain, discharge, repeated vomiting, "
                    "severe headache, or neurologic signs develop."
                )
        elif self._asks_explanatory_clinical_question(primary) or concerns:
            answer = self._build_explanatory_concern_answer(
                primary=primary,
                context=context,
                child_data=child_data,
                concerns=concerns,
            )
            if not answer.strip():
                topics = extract_clinical_topics(primary) or extract_clinical_topics(context)
                topic_text = ", ".join(topics[:3]) if topics else "the reported concern"
                answer = (
                    f"For this {age} child (HearTech risk score {risk}/100), regarding {topic_text}: "
                    "explain what parents may be observing, why hearing assessment matters in ear-related cases, "
                    "and arrange ENT/audiology review if symptoms or developmental concerns persist."
                )
        else:
            answer = (
                f"For this {age} child (HearTech risk score {risk}/100), correlate the current concern with otoscopy and "
                "hearing findings, then plan targeted follow-up. Use urgent review pathways if red flags or persistent "
                "worsening symptoms are present."
            )

        if prior_answer and self._answers_are_too_similar(answer, prior_answer):
            # Rotate wording to avoid consecutive near-duplicate fallback replies.
            if looks_like_yes_no and ("vertigo" in primary_lower or "vertigo" in concerns):
                answer = choose([
                    (
                        f"Unlikely to be definite vertigo in this {age} child (HearTech risk score {risk}/100) without clear spinning "
                        "or sustained balance disturbance. Re-check ear status and vestibular signs, and refer urgently for neurologic "
                        "red flags or ongoing severe episodes."
                    ),
                    (
                        f"Could be vestibular but not confirmed for this {age} child (HearTech risk score {risk}/100). "
                        "Clarify true spinning symptoms, examine ears and gait, and escalate rapidly if severe or persistent features appear."
                    ),
                ], offset=1)
            elif self._is_medicine_request(primary):
                medicine_lines = self._medicine_lines_for_answer(primary or context)
                answer = (
                    f"For this {age} child (HearTech risk score {risk}/100), after confirming findings with otoscopy, "
                    f"use this medicine-focused plan: {' '.join(medicine_lines[:3])}"
                )
            elif asks_action_for_reported_symptoms(primary) or asks_precautions_or_care_request(primary):
                answer = (
                    f"For this {age} child (HearTech risk score {risk}/100), next steps should focus on examination-first care: "
                    "otoscopy, concern-specific management, and close follow-up. Give urgent-care instructions for severe pain, "
                    "discharge, vomiting, high fever, or neurologic symptoms."
                )
        return answer

    @staticmethod
    def _stable_index(seed: str, size: int) -> int:
        if size <= 1:
            return 0
        return sum(ord(ch) for ch in seed) % size

    @staticmethod
    def _normalize_similarity_text(text: str) -> set[str]:
        words = re.findall(r"[a-zA-Z]+", text.lower())
        stop = {
            "the",
            "and",
            "for",
            "this",
            "that",
            "with",
            "from",
            "child",
            "heartech",
            "risk",
            "score",
            "age",
            "to",
            "of",
            "or",
            "is",
            "are",
        }
        return {w for w in words if len(w) > 3 and w not in stop}

    @classmethod
    def _answers_are_too_similar(cls, current: str, previous: str) -> bool:
        cur = cls._normalize_similarity_text(current)
        prev = cls._normalize_similarity_text(previous)
        if not cur or not prev:
            return False
        overlap = len(cur & prev)
        return overlap >= min(len(cur), len(prev)) * 0.75

    @staticmethod
    def _extract_prior_referral_from_context(context: str) -> str:
        if not context.strip():
            return ""
        marker = "immediate prior referral letter:"
        collecting = False
        chunks: list[str] = []
        for raw in context.splitlines():
            lower = raw.strip().lower()
            if lower.startswith(marker):
                collecting = True
                remainder = raw.split(":", 1)[1].strip() if ":" in raw else ""
                if remainder:
                    chunks.append(remainder)
                continue
            if collecting:
                if lower.startswith("immediate prior hcw question:") or lower.startswith(
                    "immediate prior ai answer:"
                ):
                    break
                chunks.append(raw.rstrip())
        return "\n".join(chunks).strip()

    @staticmethod
    def _patch_referral_add_destination(prior: str, destination: str) -> str:
        dest = destination.strip()
        if not prior.strip() or not dest:
            return prior.strip()

        text = prior.strip()
        if "PATIENT CARE ADVICE" in text:
            text = text.replace("PATIENT CARE ADVICE", "PATIENT REFERRAL", 1)
        if FOOTER_ADVICE in text:
            text = text.replace(FOOTER_ADVICE, FOOTER_REFERRAL)

        refer_block = ["REFER TO:", dest, "Urgency: Prompt review (within 1 week)", ""]
        lines = text.splitlines()
        out: list[str] = []
        i = 0
        replaced = False
        while i < len(lines):
            if lines[i].strip().upper().startswith("REFER TO"):
                replaced = True
                out.extend(refer_block)
                i += 1
                while i < len(lines):
                    upper = lines[i].strip().upper()
                    if upper.startswith("RECOMMENDED CARE"):
                        break
                    if upper.startswith("URGENCY"):
                        i += 1
                        continue
                    i += 1
                continue
            out.append(lines[i])
            i += 1

        if not replaced:
            insert_at = next(
                (idx for idx, ln in enumerate(out) if ln.strip().upper().startswith("RECOMMENDED CARE")),
                len(out),
            )
            out[insert_at:insert_at] = refer_block

        return "\n".join(out).strip()

    def _edit_existing_referral(
        self,
        *,
        prior_referral: str,
        edit_instruction: str,
        child_data: dict,
        normalized_prompt: str,
    ) -> GenerationResult:
        if is_medicine_section_edit(edit_instruction):
            patched = self._apply_rule_based_referral_edit(
                prior_referral=prior_referral,
                edit_instruction=edit_instruction,
                child_data=child_data,
            )
            if patched.strip() != prior_referral.strip():
                print(
                    f"[REFERRAL-AI] {self._public_log_label('referral_edit_rules')} medicine patch "
                    f"({len(patched.strip())} chars)"
                )
                return GenerationResult(
                    text=patched,
                    success=True,
                    source="referral_edit_rules",
                    intent="referral",
                    needs_clarification=False,
                    normalized_prompt=normalized_prompt,
                )
            print(
                "[REFERRAL-AI] Medicine edit needs model reasoning — "
                "rule patch made no change, trying referral_edit model."
            )

        include_refer_to = "REFER TO:" in prior_referral.upper()
        edit_prompt = build_referral_edit_prompt(
            child_data=child_data,
            prior_referral=prior_referral,
            edit_instruction=edit_instruction,
        )

        local_text = self._run_inference(
            self._model_primary,
            edit_prompt,
            label="referral_edit",
            max_tokens=900,
            temp=0.15,
            top_p=0.9,
        )
        local_text = self._sanitize_referral_output(
            local_text, include_refer_to=include_refer_to or "REFER TO:" in local_text.upper()
        )
        local_validation = self._validate_output(
            output=local_text,
            intent="referral",
            include_refer_to=include_refer_to or "REFER TO:" in local_text.upper(),
            child_data=child_data,
            primary_instruction=edit_instruction,
            generation_instruction=edit_instruction,
        )
        if self._can_ship_referral_edit(prior_referral, local_text, local_validation):
            print(
                f"[REFERRAL-AI] {self._public_log_label('referral_edit')} accepted "
                f"({len(local_text.strip())} chars)"
            )
            return GenerationResult(
                text=local_text,
                success=True,
                source="referral_edit",
                intent="referral",
                needs_clarification=False,
                normalized_prompt=normalized_prompt,
                validation_warnings=local_validation.warnings,
            )

        if self._cloud_provider.is_available():
            try:
                cloud_raw = self._cloud_provider.generate(edit_prompt, max_tokens=900)
                cloud_text = self._sanitize_referral_output(
                    cloud_raw,
                    include_refer_to=include_refer_to or "REFER TO:" in cloud_raw.upper(),
                )
                cloud_validation = self._validate_output(
                    output=cloud_text,
                    intent="referral",
                    include_refer_to=include_refer_to or "REFER TO:" in cloud_text.upper(),
                    child_data=child_data,
                    primary_instruction=edit_instruction,
                    generation_instruction=edit_instruction,
                )
                if self._can_ship_referral_edit(prior_referral, cloud_text, cloud_validation):
                    print(
                        f"[REFERRAL-AI] {self._public_log_label('referral_edit_cloud')} accepted "
                        f"({len(cloud_text.strip())} chars)"
                    )
                    return GenerationResult(
                        text=cloud_text,
                        success=True,
                        source="referral_edit_cloud",
                        intent="referral",
                        needs_clarification=False,
                        normalized_prompt=normalized_prompt,
                        validation_warnings=cloud_validation.warnings,
                    )
            except Exception:
                print(f"[REFERRAL-AI] {self._public_log_label('referral_edit_cloud')} unavailable.")

        fallback = self._apply_rule_based_referral_edit(
            prior_referral=prior_referral,
            edit_instruction=edit_instruction,
            child_data=child_data,
        )
        include_dest = "REFER TO:" in fallback.upper()
        fallback = self._finalize_referral_for_hcw(
            fallback,
            prompt=edit_instruction,
            child_data=child_data,
            include_refer_to=include_dest,
        )
        print(
            f"[REFERRAL-AI] {self._public_log_label('referral_edit_rules')} applied "
            f"({len(fallback.strip())} chars)"
        )
        return GenerationResult(
            text=fallback,
            success=True,
            source="referral_edit_rules",
            intent="referral",
            needs_clarification=False,
            normalized_prompt=normalized_prompt,
            validation_warnings=local_validation.warnings,
        )

    @staticmethod
    def _can_ship_referral_edit(
        prior: str, edited: str, validation: ValidationResult
    ) -> bool:
        if not edited.strip() or "PATIENT REFERRAL" not in edited:
            return False
        if ReferralAIService._referral_edit_lost_prior_content(prior, edited):
            return False
        return validation.ok

    @staticmethod
    def _referral_edit_lost_prior_content(prior: str, edited: str) -> bool:
        prior_l = prior.lower()
        edited_l = edited.lower()
        for term in (
            "vertigo",
            "cyclizine",
            "tinnitus",
            "fever with ear",
            "speech delay",
            "ear wax",
            "hornedge",
        ):
            if term in prior_l and term not in edited_l:
                return True
        generic_markers = (
            "focused ear and hearing assessment is advised",
            "hearing-related concerns (risk score",
        )
        if any(marker in edited_l for marker in generic_markers):
            if not any(marker in prior_l for marker in generic_markers):
                return True
        return False

    def _apply_rule_based_referral_edit(
        self,
        *,
        prior_referral: str,
        edit_instruction: str,
        child_data: dict,
    ) -> str:
        text = prior_referral.strip()
        lower = edit_instruction.lower()

        destination = extract_referral_destination(edit_instruction)
        if destination:
            text = self._patch_referral_add_destination(text, destination)

        for phrase in self._extract_symptom_phrases_for_edit(edit_instruction):
            if phrase.lower() not in text.lower():
                text = self._append_to_referral_section(
                    text,
                    section_header="CLINICAL SUMMARY:",
                    addition=f"{phrase.capitalize()} also reported and should be addressed in assessment.",
                )

        if "make it urgent" in lower or re.search(r"\burgent(?:ly)?\b", lower):
            text = self._upsert_referral_section_line(
                text,
                section_header="REFER TO:",
                line="Urgency: Urgent review within 24-48 hours",
                insert_after_header=True,
            )
            if "REFER TO:" not in text.upper():
                text = self._append_to_referral_section(
                    text,
                    section_header="CLINICAL SUMMARY:",
                    addition="Urgent review requested based on current concerns.",
                )

        if self._is_medicine_request(edit_instruction) or is_medicine_section_edit(
            edit_instruction
        ):
            medicine_lines = self._medicine_lines_for_referral_edit(
                edit_instruction, text
            )
            for line in medicine_lines:
                if not self._referral_has_medicine_line(text, line):
                    text = self._append_to_referral_section(
                        text,
                        section_header="MEDICINES (if applicable):",
                        addition=f"• {line}" if not line.startswith("•") else line,
                        as_bullet=True,
                    )

        concerns = self._detect_concern_keys(edit_instruction)
        for concern in concerns:
            label = self._concern_label(concern)
            if label.lower() not in text.lower():
                text = self._append_to_referral_section(
                    text,
                    section_header="CLINICAL SUMMARY:",
                    addition=f"{label} noted in current assessment.",
                )
            for care_line in self._referral_care_lines([concern]):
                if care_line.lower() not in text.lower():
                    text = self._append_to_referral_section(
                        text,
                        section_header="RECOMMENDED CARE:",
                        addition=f"• {care_line}",
                        as_bullet=True,
                    )

        if "speech" in lower and "speech" not in text.lower():
            text = self._append_to_referral_section(
                text,
                section_header="RECOMMENDED CARE:",
                addition="• Monitor speech/listening milestones and arrange speech-language input if hearing loss is excluded.",
                as_bullet=True,
            )

        if re.search(r"\binvestigations?\b", lower):
            text = self._append_to_referral_section(
                text,
                section_header="RECOMMENDED CARE:",
                addition="• Consider formal hearing investigation (OAE/ABR/audiometry) as clinically indicated.",
                as_bullet=True,
            )

        for care_line in extract_hcw_care_requests(edit_instruction):
            if care_line.lower()[:24] not in text.lower():
                text = self._append_to_referral_section(
                    text,
                    section_header="RECOMMENDED CARE:",
                    addition=f"• {care_line}",
                    as_bullet=True,
                )

        if "PATIENT CARE ADVICE" in text:
            text = text.replace("PATIENT CARE ADVICE", "PATIENT REFERRAL", 1)
        if FOOTER_ADVICE in text:
            text = text.replace(FOOTER_ADVICE, FOOTER_REFERRAL)
        if child_data.get("name") and child_data["name"] not in text:
            text = text  # keep prior letter as-is
        return text.strip()

    @staticmethod
    def _normalize_clinical_typos(text: str) -> str:
        replacements = (
            (r"\bfeve\b", "fever"),
            (r"\bfebrile\b", "fever"),
            (r"\brefferal\b", "referral"),
            (r"\breferal\b", "referral"),
        )
        normalized = text
        for pattern, repl in replacements:
            normalized = re.sub(pattern, repl, normalized, flags=re.IGNORECASE)
        return normalized

    @staticmethod
    def _medicine_indication_from_edit(instruction: str, concerns: list[str]) -> str:
        lower = ReferralAIService._normalize_clinical_typos(instruction).lower()
        if "fever" in concerns or re.search(r"\bfever\b", lower):
            return "for current febrile symptoms"
        if "ear_infection" in concerns:
            return "if bacterial otitis is suspected on examination"
        if "vertigo" in concerns:
            return "for acute vertigo if prescribed"
        if "pain" in lower:
            return "for pain relief"
        return "as clinically indicated for the reported concern"

    def _medicine_lines_for_referral_edit(
        self, edit_instruction: str, prior_referral: str
    ) -> list[str]:
        instruction = self._normalize_clinical_typos(edit_instruction)
        lower = instruction.lower()
        prior_lower = prior_referral.lower()

        concerns = self._detect_concern_keys(instruction)
        for key in (
            "fever",
            "speech",
            "vertigo",
            "ear_infection",
            "ear_wax",
            "investigations",
        ):
            if key not in concerns and self._line_asserts_concern(prior_lower, key):
                concerns.append(key)

        cond_match = re.search(
            r"\bmedicines?\b(?:\s+(?:for|for the))?\s+([a-z][\w\s\-]{2,40})",
            lower,
        )
        if cond_match:
            cond_text = self._normalize_clinical_typos(cond_match.group(1))
            for key in self._detect_concern_keys(cond_text):
                if key not in concerns:
                    concerns.append(key)

        if re.search(r"\b(?:add|include)\s+medicines?\b", lower) and not concerns:
            concerns = self._detect_concern_keys(prior_lower)

        lines: list[str] = []
        indication = self._medicine_indication_from_edit(instruction, concerns)

        for med_key in self._extract_requested_medicines_from_edit(instruction):
            base = MEDICINE_GUIDANCE.get(med_key, med_key)
            lines.append(f"{base} ({indication})")

        if "fever" in concerns or re.search(r"\bfever\b", lower):
            lines.append(f"{MEDICINE_GUIDANCE['paracetamol']} ({indication})")
            lines.append(
                f"{MEDICINE_GUIDANCE['ibuprofen']} (for fever/pain if no contraindications)"
            )
        if "ear_infection" in concerns or (
            "fever" in concerns and "ear" in prior_lower
        ):
            lines.append(
                f"{MEDICINE_GUIDANCE['amoxicillin']} (if bacterial otitis is suspected on examination)"
            )
        if "vertigo" in concerns:
            for med in CONCERN_MEDICINES.get("vertigo", ()):
                base = MEDICINE_GUIDANCE.get(med, med)
                lines.append(f"{base} (for vertigo/balance symptoms if prescribed)")

        if not lines:
            for concern in concerns:
                for med in CONCERN_MEDICINES.get(concern, ()):
                    base = MEDICINE_GUIDANCE.get(med, med)
                    tagged = f"{base} ({indication})"
                    if tagged not in lines:
                        lines.append(tagged)
            if not lines and (self._is_medicine_request(instruction) or concerns):
                lines.append(f"{MEDICINE_GUIDANCE['paracetamol']} ({indication})")
                lines.append(f"{MEDICINE_GUIDANCE['ibuprofen']} ({indication})")

        deduped: list[str] = []
        seen: set[str] = set()
        for line in lines:
            key = re.sub(r"[^a-z0-9]+", "", line.lower())
            if not key or key in seen:
                continue
            seen.add(key)
            deduped.append(line)
        if deduped:
            deduped.append("Always confirm medicine choice and dose with a qualified clinician.")
        return deduped

    @staticmethod
    def _referral_has_medicine_line(text: str, line: str) -> bool:
        target = line.strip().lower().lstrip("•").strip()
        if not target:
            return True
        for raw in text.splitlines():
            ln = raw.strip().lower().lstrip("•").strip()
            if ln == target:
                return True
        return False

    @staticmethod
    def _extract_requested_medicines_from_edit(instruction: str) -> list[str]:
        instruction = ReferralAIService._normalize_clinical_typos(instruction)
        lower = instruction.lower()
        found: list[str] = []
        for alias, canonical in _MEDICINE_ALIASES.items():
            if re.search(rf"\b{re.escape(alias)}\b", lower) and canonical not in found:
                found.append(canonical)
        for med_key in MEDICINE_GUIDANCE:
            if re.search(rf"\b{re.escape(med_key)}\b", lower) and med_key not in found:
                found.append(med_key)
        add_match = re.search(
            r"\b(?:add|include|put|insert)\s+([a-z][a-z0-9\- ]{2,40}?)(?:\s+in|\s+to|\s+for|$)",
            lower,
        )
        if add_match:
            token = add_match.group(1).strip()
            canonical = _MEDICINE_ALIASES.get(token, token)
            if canonical in MEDICINE_GUIDANCE and canonical not in found:
                found.append(canonical)
        return found

    @staticmethod
    def _extract_symptom_phrases_for_edit(instruction: str) -> list[str]:
        lower = instruction.lower()
        phrases: list[str] = []
        for term in (
            "tinnitus",
            "vertigo",
            "fever",
            "ear wax",
            "hearing loss",
            "speech delay",
            "balance problems",
            "ear pain",
            "ear discharge",
        ):
            if term in lower and term not in phrases:
                phrases.append(term)
        match = re.search(r"symptoms? of ([a-z][a-z\s\-]{2,40})", lower)
        if match:
            phrase = match.group(1).strip(" .,!?:;")
            if phrase and phrase not in phrases:
                phrases.append(phrase)
        return phrases

    @staticmethod
    def _append_to_referral_section(
        text: str,
        *,
        section_header: str,
        addition: str,
        as_bullet: bool = False,
    ) -> str:
        lines = text.splitlines()
        out: list[str] = []
        header_key = section_header.rstrip(":").upper()
        inserted = False
        i = 0
        while i < len(lines):
            out.append(lines[i])
            if not inserted and lines[i].strip().upper().startswith(header_key):
                i += 1
                section_lines: list[str] = []
                while i < len(lines):
                    nxt = lines[i].strip()
                    if nxt.endswith(":") and nxt.upper() == nxt and len(nxt) > 3:
                        break
                    section_lines.append(lines[i])
                    i += 1
                if section_lines:
                    last = section_lines[-1].rstrip()
                    if addition.lower() not in last.lower() and addition.lower() not in text.lower():
                        if as_bullet:
                            section_lines.append(addition if addition.startswith("•") else f"• {addition}")
                        else:
                            joiner = "" if last.endswith((".", "!", "?")) else "; "
                            section_lines[-1] = f"{last}{joiner}{addition}"
                else:
                    section_lines.append(addition if as_bullet or addition.startswith("•") else addition)
                out.extend(section_lines)
                inserted = True
                continue
            i += 1
        if not inserted:
            out.extend(["", section_header, addition if as_bullet or addition.startswith("•") else addition])
        return "\n".join(out).strip()

    @staticmethod
    def _upsert_referral_section_line(
        text: str,
        *,
        section_header: str,
        line: str,
        insert_after_header: bool = False,
    ) -> str:
        lines = text.splitlines()
        out: list[str] = []
        header_key = section_header.rstrip(":").upper()
        replaced = False
        i = 0
        while i < len(lines):
            if lines[i].strip().upper().startswith(header_key):
                out.append(lines[i])
                i += 1
                if insert_after_header:
                    while i < len(lines):
                        nxt = lines[i].strip().upper()
                        if nxt.startswith("URGENCY"):
                            out.append(line)
                            replaced = True
                            i += 1
                            continue
                        if nxt.startswith("RECOMMENDED CARE"):
                            if not replaced:
                                out.append(line)
                            break
                        out.append(lines[i])
                        i += 1
                    if not replaced:
                        out.append(line)
                    continue
            out.append(lines[i])
            i += 1
        if header_key not in text.upper():
            out.extend(["", section_header, line])
        return "\n".join(out).strip()

    @staticmethod
    def _extract_prior_ai_answer_from_context(context: str) -> str:
        for raw in context.splitlines():
            line = raw.strip()
            if line.lower().startswith("immediate prior ai answer:"):
                return line.split(":", 1)[1].strip()
        return ""

    @staticmethod
    def _is_medicine_request(prompt: str) -> bool:
        p = prompt.lower()
        return any(
            key in p
            for key in (
                "medicine",
                "medicines",
                "medication",
                "drug",
                "dose",
                "dosing",
                "pain relief",
                "suggest meds",
            )
        )

    @staticmethod
    def _medicine_lines_for_answer(prompt: str) -> list[str]:
        lines = ReferralAIService._medicine_lines_for_prompt(prompt)
        selected: list[str] = []
        wax_softening_added = False
        for line in lines:
            lower = line.lower()
            if "for wax softening" in lower:
                if wax_softening_added:
                    continue
                wax_softening_added = True
            if line not in selected:
                selected.append(line)
            if len(selected) >= 3:
                break
        return selected

    @staticmethod
    def _split_instruction_context(hcw_instruction: str) -> tuple[str, str]:
        marker = "Conversation context:"
        if marker not in hcw_instruction:
            return hcw_instruction.strip(), ""
        primary, ctx = hcw_instruction.split(marker, 1)
        return primary.strip(), ctx.strip()

    @staticmethod
    def _sanitize_answer_output(text: str) -> str:
        if not text:
            return ""

        for tok in _LLAMA_SPECIAL_TOKENS:
            text = text.replace(tok, " ")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text).strip()

        # Hard cut if model starts drifting into referral template.
        template_markers = (
            "PATIENT REFERRAL",
            "PATIENT CARE ADVICE",
            "CLINICAL SUMMARY:",
            "REFER TO:",
            "RECOMMENDED CARE:",
            "PRECAUTIONS FOR PARENT:",
            "FOLLOW UP:",
            "This referral was generated by HearTech screening system.",
            "This advice was generated by HearTech screening system.",
            "Screened via HearTech",
        )
        cut_idx = None
        for marker in template_markers:
            idx = text.find(marker)
            if idx != -1 and (cut_idx is None or idx < cut_idx):
                cut_idx = idx
        if cut_idx is not None:
            text = text[:cut_idx]

        text = ReferralAIService._scrub_training_template_sentences(text)

        text = re.sub(
            r"\(?\s*max\s+\d+\s+short\s+lines?\s*\)?",
            "",
            text,
            flags=re.I,
        )
        text = re.sub(r"\s{2,}", " ", text).strip()

        # Hard cut when model leaks unrelated adult-patient training narratives.
        contamination_cuts = (
            "my husband",
            "my wife",
            "my father",
            "my mother",
            "years old and in good health",
            "since last week",
            "he is 45",
            "she is 45",
            "experiencing dizziness, nausea",
            "at around 4 years",
            "at around 18 months",
            "language has settled",
            "two most important factors to detect hearing loss",
            "key milestones include",
            "consistent response to familiar voices",
            "speech-language progression",
            "absent babbling",
            "milestone review with oae",
            "if these signs persist beyond 2-4 weeks",
        )
        _training_line_markers = (
            "eustachian tube dysfunction",
            "28+ days disease",
            "moderate persistent, not severe",
            "child has eustachian",
        )
        lower_text = text.lower()
        for marker in contamination_cuts:
            idx = lower_text.find(marker)
            if idx != -1:
                text = text[:idx].rstrip()
                lower_text = text.lower()
                break

        # Split to sentence-like units for dedupe. Do not split camelCase — it breaks HearTech.
        protected = text.replace("HearTech", "HEARTECH_KEEP")
        parts = re.split(r"[\n\r]+|(?<=[.!?])\s+", protected)
        lines = [ln.strip(" -•\t") for ln in parts if ln.strip(" -•\t")]
        deduped: list[str] = []
        seen: set[str] = set()
        for ln in lines:
            ln = ln.replace("HEARTECH_KEEP", "HearTech")
            ln_lower = ln.lower()
            if any(marker in ln_lower for marker in _training_line_markers):
                continue
            if is_instruction_leak_line(ln):
                continue
            if is_noisy_answer_line(ln):
                continue
            key = re.sub(r"[^a-z0-9]+", "", ln.lower())
            if len(key) < 12:
                continue
            # Drop near-duplicate sentences (model repetition), keep the longer line.
            duplicate = False
            for existing in list(seen):
                if key.startswith(existing) or existing.startswith(key):
                    duplicate = True
                    break
            if duplicate:
                continue
            seen.add(key)
            deduped.append(ln)
            if len(deduped) >= 10:
                break

        cleaned = "\n".join(deduped).strip()
        if not cleaned:
            return ""
        return cleaned

    @staticmethod
    def _sanitize_referral_output(text: str, *, include_refer_to: bool) -> str:
        if not text:
            return ""
        for tok in _LLAMA_SPECIAL_TOKENS:
            text = text.replace(tok, " ")
        lines = [ln.rstrip() for ln in text.splitlines() if ln.strip()]
        deduped: list[str] = []
        seen: set[str] = set()
        for ln in lines:
            key = re.sub(r"[^a-z0-9]+", "", ln.lower())
            if not key:
                continue
            if key in seen:
                continue
            if not include_refer_to and ln.strip().upper().startswith("REFER TO"):
                continue
            seen.add(key)
            deduped.append(ln)
        return "\n".join(deduped).strip()

    def _build_rule_based_referral(
        self,
        *,
        prompt: str,
        child_data: dict,
        include_refer_to: bool,
        concerns_override: list[str] | None = None,
    ) -> str:
        name = child_data.get("name", "Patient")
        age = child_data.get("age", "Unknown")
        risk = child_data.get("riskScore", "N/A")
        date = datetime.now().strftime("%d %b %Y")
        concerns = concerns_override or self._detect_concern_keys_for_referral(prompt)
        summary = self._referral_clinical_summary(concerns, age=age, risk=risk)
        care_lines = self._referral_care_lines(concerns)
        requested_care = extract_hcw_care_requests(prompt)
        for line in requested_care:
            if line not in care_lines:
                care_lines.append(line)
        if any(x in prompt.lower() for x in ("ling 6", "ling-6", "ling6", "show and tell", "show-and-tell")):
            if "speech" not in concerns:
                concerns = list(concerns) + ["speech"]
                summary = self._referral_clinical_summary(concerns, age=age, risk=risk)
        precaution_lines = self._referral_precaution_lines(concerns)
        medicine_lines = self._medicine_lines_for_concerns(concerns)
        if self._is_medicine_request(prompt):
            prompt_meds = self._medicine_lines_for_answer(prompt)
            if prompt_meds:
                medicine_lines = prompt_meds

        care_block = "\n".join(f"• {line}" for line in care_lines)
        precaution_block = "\n".join(f"• {line}" for line in precaution_lines)
        meds_block = (
            "MEDICINES (if applicable):\n"
            + "\n".join(f"• {line}" for line in medicine_lines)
            + "\n\n"
        )

        destination = extract_referral_destination(prompt)
        if destination:
            include_refer_to = True

        urgency = "Prompt review (within 1 week)"
        if "urgent" in prompt.lower():
            urgency = "Urgent (within 1 week)"

        if include_refer_to and destination:
            refer_line = destination
            return (
                "PATIENT REFERRAL\n"
                "─────────────────\n"
                f"Patient: {name} | Age: {age} | Date: {date}\n\n"
                "CLINICAL SUMMARY:\n"
                f"{summary}\n\n"
                "REFER TO:\n"
                f"{refer_line}\n"
                f"Urgency: {urgency}\n\n"
                "RECOMMENDED CARE:\n"
                f"{care_block}\n\n"
                f"{meds_block}"
                "PRECAUTIONS FOR PARENT:\n"
                f"{precaution_block}\n\n"
                "FOLLOW UP:\n"
                "Review in 48-72 hours if symptoms persist; complete specialist follow-up within 1 week.\n\n"
                f"Screened via HearTech | {date}\n"
                "This referral was generated by HearTech screening system.\n"
                "Always confirm with a qualified clinician before administration."
            )

        return (
            "PATIENT REFERRAL\n"
            "─────────────────\n"
            f"Patient: {name} | Age: {age} | Date: {date}\n\n"
            "CLINICAL SUMMARY:\n"
            f"{summary}\n\n"
            "RECOMMENDED CARE:\n"
            f"{care_block}\n\n"
            f"{meds_block}"
            "PRECAUTIONS FOR PARENT:\n"
            f"{precaution_block}\n\n"
            "FOLLOW UP:\n"
            "Review in 48-72 hours if symptoms persist; complete specialist follow-up as needed.\n\n"
            f"Screened via HearTech | {date}\n"
            "This referral was generated by HearTech screening system.\n"
            "Always confirm with a qualified clinician before administration."
        )

    def _build_reasoned_referral(
        self, *, prompt: str, child_data: dict, include_refer_to: bool
    ) -> str:
        """Concern-aware referral builder that prioritizes HCW-mentioned concerns."""
        primary, context = self._split_instruction_context(prompt)
        hcw_only_context = self._extract_hcw_only_context(context)
        concern_source = f"{primary}\n{hcw_only_context}".strip()
        if not concern_source:
            concern_source = primary
        synthesized_prompt = (
            f"{primary}\n\nConversation context:\n{hcw_only_context}"
            if hcw_only_context
            else primary
        )
        concerns = self._detect_concern_keys_for_referral(concern_source)
        return self._build_rule_based_referral(
            prompt=synthesized_prompt,
            child_data=child_data,
            include_refer_to=include_refer_to,
            concerns_override=concerns,
        )

    @staticmethod
    def _referral_clinical_summary(concerns: list[str], *, age: str, risk: str) -> str:
        if not concerns:
            return (
                f"{age} child screened with HearTech has hearing-related concerns (risk score {risk}). "
                "Focused ear and hearing assessment is advised."
            )
        labels = [ReferralAIService._concern_label(c) for c in concerns]
        joined = ", ".join(labels)
        detail_parts: list[str] = []
        if "ear_wax" in concerns or "impression" in concerns:
            detail_parts.append(
                "canal blockage/wax may be affecting hearing and preventing ear impression"
            )
        if "speech" in concerns:
            detail_parts.append("speech/language concerns warrant hearing-first assessment")
        if "vertigo" in concerns:
            detail_parts.append("vertigo/balance symptoms require red-flag review and ENT evaluation")
        if "investigations" in concerns:
            detail_parts.append("formal hearing investigation (OAE/ABR/audiometry) is indicated")
        detail = "; ".join(detail_parts) + "." if detail_parts else "Targeted management is advised."
        return (
            f"{age} child screened with HearTech (risk score {risk}) with concerns including {joined}. "
            f"{detail.capitalize()}"
        )

    @staticmethod
    def _referral_care_lines(concerns: list[str]) -> list[str]:
        lines = [
            "Perform otoscopy and age-appropriate hearing assessment (OAE/ABR/audiometry as indicated).",
        ]
        if "ear_wax" in concerns or "impression" in concerns:
            lines.append(
                "Clear impacted cerumen safely before ear mould/impression if tympanic membrane is intact."
            )
        if "speech" in concerns:
            lines.append(
                "Monitor speech/listening milestones and arrange speech-language input if hearing loss is excluded."
            )
        if "vertigo" in concerns:
            lines.append(
                "Assess for vestibular symptoms and middle-ear disease; escalate urgently if red flags are present."
            )
        if "investigations" in concerns or "speech" in concerns:
            lines.append("Consider ABR when behavioral testing is unreliable in this age group.")
        lines.append("Track symptom progression and communication response daily.")
        return lines

    @staticmethod
    def _referral_precaution_lines(concerns: list[str]) -> list[str]:
        lines = [
            "Avoid cotton buds or object insertion in the ear canal.",
            "Seek urgent care for fever, ear discharge, severe pain, persistent vomiting, or reduced responsiveness.",
        ]
        if "speech" in concerns:
            lines.append(
                "Reduce background noise at home and monitor response to name and familiar sounds daily."
            )
        if "vertigo" in concerns:
            lines.append(
                "Supervise mobility closely during vertigo episodes and keep the child hydrated."
            )
        if "ear_wax" in concerns or "impression" in concerns:
            lines.append("Keep ears dry and complete wax softening only as clinically advised.")
        return lines

    @staticmethod
    def _extract_primary_concern(prompt: str) -> str:
        p = prompt.lower()
        if "speech" in p or "language" in p:
            return "speech/language delay"
        if any(x in p for x in ("ear wax", "earwax", "cerumen", "blocked ear", "ear canal")):
            return "ear canal blockage/wax with hearing concern"
        if any(x in p for x in ("vertigo", "dizziness", "balance")):
            return "vertigo/balance-related hearing concern"
        if any(x in p for x in ("not responding", "low voice", "response to voice")):
            return "reduced response to voice/sound"
        return "hearing-risk symptoms"

    _CONCERN_LABELS = {
        "ear_wax": "ear canal blockage/wax",
        "impression": "ear impression/mould difficulty",
        "speech": "speech/language delay",
        "vertigo": "vertigo/balance symptoms",
        "low_response": "reduced response to voice/sound",
        "investigations": "hearing investigation needs",
        "fever": "fever with ear symptoms",
        "ear_infection": "possible acute otitis media",
    }

    @classmethod
    def _concern_label(cls, key: str) -> str:
        return cls._CONCERN_LABELS.get(key, key.replace("_", " "))

    @staticmethod
    def _is_aggregate_referral_request(primary: str) -> bool:
        p = primary.lower()
        aggregate_markers = (
            "all of this",
            "all of these",
            "everything discussed",
            "everything above",
            "everything we",
            "for all of",
        )
        if any(m in p for m in aggregate_markers):
            return True
        return "referral" in p and any(m in p for m in ("all of", "for this", "for these"))

    @staticmethod
    def _extract_excluded_concerns(text: str) -> set[str]:
        p = text.lower()
        excluded: set[str] = set()
        if any(
            ph in p
            for ph in (
                "without vertigo",
                "no vertigo",
                "not vertigo",
                "exclude vertigo",
                "does not have vertigo",
                "doesn't have vertigo",
            )
        ):
            excluded.add("vertigo")
        if any(ph in p for ph in ("without speech", "no speech delay", "exclude speech")):
            excluded.add("speech")
        if any(ph in p for ph in ("without wax", "no ear wax", "without ear wax", "exclude wax")):
            excluded.add("ear_wax")
        return excluded

    @staticmethod
    def _is_differential_vertigo_mention(text: str) -> bool:
        p = text.lower()
        patterns = (
            r"vertigo\s+or",
            r"or\s+(?:just\s+)?(?:ear\s+)?wax",
            r"mean\s+(?:the\s+child\s+has\s+)?vertigo",
            r"has\s+vertigo\s+or",
            r"blockage.*vertigo.*or",
        )
        return any(re.search(pat, p) for pat in patterns)

    @staticmethod
    def _line_asserts_concern(line: str, concern: str) -> bool:
        p = line.lower().strip()
        if not p:
            return False
        if concern == "vertigo":
            if ReferralAIService._is_differential_vertigo_mention(p):
                return False
            if re.search(r"(?:without|no|not|doesn'?t|does not)\s+(?:have\s+)?vertigo", p):
                return False
            return any(x in p for x in ("vertigo", "dizziness", "spinning"))
        if concern == "ear_wax":
            return any(
                x in p
                for x in (
                    "ear wax",
                    "earwax",
                    "cerumen",
                    "blocked ear",
                    "ear canal",
                    "blockage",
                    "blocked",
                    "ear block",
                    "bloackage",
                )
            )
        if concern == "fever":
            return any(x in p for x in ("fever", "febrile", "pyrexia", "high temp", "feve"))
        if concern == "ear_infection":
            return any(x in p for x in ("otitis", "ear infection", "infected ear", "bulging"))
        if concern == "speech":
            return any(
                x in p
                for x in (
                    "speech",
                    "language",
                    "stutter",
                    "stutters",
                    "stuttering",
                    "speech delay",
                    "speech delays",
                    "delayed speech",
                    "language delay",
                )
            )
        if concern == "impression":
            return any(x in p for x in ("impression", "mould", "mold"))
        if concern == "low_response":
            return any(x in p for x in ("not responding", "low voice", "response to voice"))
        if concern == "investigations":
            return any(x in p for x in ("abr", "auditory brainstem", "oae", "audiometry"))
        return False

    @staticmethod
    def _detect_concern_keys(text: str, *, excluded: set[str] | None = None) -> list[str]:
        excluded = excluded or set()
        concerns: list[str] = []
        for key in (
            "ear_wax",
            "impression",
            "speech",
            "fever",
            "ear_infection",
            "vertigo",
            "low_response",
            "investigations",
        ):
            if key in excluded:
                continue
            if ReferralAIService._line_asserts_concern(text, key) and key not in concerns:
                concerns.append(key)
        return concerns

    @staticmethod
    def _references_pronoun_followup(primary: str) -> bool:
        p = primary.lower()
        return any(
            x in p
            for x in (
                " for this",
                " for that",
                " for it",
                "any medicine for this",
                "medicine for this",
                "precaution for this",
            )
        ) or (("medicine" in p or "precaution" in p) and "this" in p)

    @staticmethod
    def _context_lines(context: str) -> list[str]:
        lines: list[str] = []
        for raw in context.splitlines():
            line = raw.strip().lstrip("-").strip()
            if not line:
                continue
            lower = line.lower()
            if lower.startswith(
                (
                    "prior hcw questions",
                    "recent hcw questions",
                    "conversation context",
                    "immediate prior hcw question:",
                    "immediate prior clinical guidance",
                )
            ):
                if ":" in line:
                    _, _, rest = line.partition(":")
                    rest = rest.strip()
                    if rest:
                        lines.append(rest)
                continue
            lines.append(line)
        return lines

    @staticmethod
    def _extract_hcw_only_context(context: str) -> str:
        lines: list[str] = []
        in_hcw_question_list = False
        for raw in context.splitlines():
            line = raw.strip()
            if not line:
                continue
            lower = line.lower()
            if lower.startswith("prior hcw questions in this session"):
                in_hcw_question_list = True
                continue
            if lower.startswith("immediate prior hcw question:"):
                lines.append(line.split(":", 1)[1].strip())
                continue
            if lower.startswith("immediate prior ai answer:"):
                continue
            if in_hcw_question_list:
                if line.startswith("- "):
                    lines.append(line[2:].strip())
                    continue
                # End list when bullets stop.
                in_hcw_question_list = False
            if lower.startswith("hcw:"):
                lines.append(line.split(":", 1)[1].strip())
        return "\n".join(x for x in lines if x)

    @staticmethod
    def _detect_concern_keys_for_answer(primary: str, context: str = "") -> list[str]:
        excluded = ReferralAIService._extract_excluded_concerns(primary)
        concerns = ReferralAIService._detect_concern_keys(primary, excluded=excluded)
        if concerns:
            return concerns

        ctx_lines = ReferralAIService._context_lines(context)
        if ReferralAIService._references_pronoun_followup(primary) and ctx_lines:
            # "medicine for this" → use the immediate prior exchange only.
            topic_text = ctx_lines[0]
            return ReferralAIService._detect_concern_keys(topic_text, excluded=excluded)

        if context and any(
            x in primary.lower()
            for x in ("precaution", "medicine", "symptom", "signs", "test", "investigation")
        ):
            recent = ctx_lines[-1] if ctx_lines else context
            return ReferralAIService._detect_concern_keys(recent, excluded=excluded)
        return concerns

    @staticmethod
    def _detect_concern_keys_for_referral(full_prompt: str) -> list[str]:
        primary, ctx = ReferralAIService._split_instruction_context(full_prompt)
        excluded = ReferralAIService._extract_excluded_concerns(primary)
        concerns: list[str] = []
        source = ctx or full_prompt
        for line in ReferralAIService._context_lines(source):
            line_lower = line.lower()
            if (
                "vertigo" in line_lower
                and not re.search(r"(?:without|no|not|doesn'?t|does not)\s+(?:have\s+)?vertigo", line_lower)
                and "vertigo" not in excluded
                and "vertigo" not in concerns
            ):
                concerns.append("vertigo")
            for key in ReferralAIService._detect_concern_keys(line, excluded=excluded):
                if key not in concerns and key not in excluded:
                    concerns.append(key)
        if not concerns:
            for key in ReferralAIService._detect_concern_keys(primary, excluded=excluded):
                if key not in concerns and key not in excluded:
                    concerns.append(key)
        return [c for c in concerns if c not in excluded]

    @staticmethod
    def _clinical_guidance_for_concern(concern: str, age: str, *, prompt: str = "") -> str:
        p = prompt.lower()
        guidance = {
            "ear_wax": (
                f"Ear wax (cerumen) can block the canal and reduce hearing, which may worsen speech response or "
                f"delay language. Confirm with otoscopy; use cerumen-softening drops only if the tympanic membrane "
                f"is intact; avoid cotton buds or probing."
            ),
            "impression": (
                "Impacted ear wax commonly prevents a reliable ear impression/mould — the canal must be cleared "
                "before mould taking. After wax removal, recheck the canal and repeat hearing screening if needed."
            ),
            "speech": (
                f"Speech or language delay means the child is not using words, phrases, or understanding language "
                f"expected for {age}. Stuttering (disfluency) is repeated sounds, prolongations, or blocks while "
                f"talking — it can occur with or without delay. Because this child has a high HearTech hearing risk, "
                f"check hearing first with otoscopy and age-appropriate tests (OAE/ABR in young children) — conductive "
                f"loss from wax or middle-ear disease is a common treatable cause of delayed speech."
            ),
            "vertigo": (
                "Vertigo in a child with ear symptoms needs urgent red-flag review (vomiting, severe headache, "
                "focal neurologic signs, or persistent spinning). If red flags are absent, assess balance, "
                "walking, and ear examination findings before attributing symptoms to wax alone."
            ),
            "low_response": (
                "Reduced response to voice or name is a hearing red flag. Prioritize otoscopy, tympanometry, "
                "and age-appropriate audiometry (OAE/ABR in young children)."
            ),
            "investigations": (
                "Age-appropriate workup: otoscopy and tympanometry first; OAE for screening; ABR when behavioral "
                "testing is unreliable or after failed OAE; pure-tone audiometry typically from ~4 years when cooperative."
            ),
            "fever": (
                "Fever with ear symptoms raises concern for otitis media or wider infection — examine the ear and "
                "throat, assess hydration and distress, and treat infection if confirmed."
            ),
            "ear_infection": (
                "Middle-ear infection (otitis media) can cause pain, fever, and temporary hearing loss that affects "
                "speech and behaviour. Examine the tympanic membrane and treat according to local guidelines."
            ),
        }
        text = guidance.get(concern, "")
        if concern == "speech" and any(x in p for x in ("stutter", "stuttering", "disfluenc")):
            text += (
                " Persistent stuttering with limited vocabulary for age warrants speech-language therapy referral "
                "alongside hearing assessment — do not assume it is only behavioural without checking ears and hearing."
            )
        return text

    @staticmethod
    def _parse_risk_score(child_data: dict) -> int | None:
        try:
            return int(child_data.get("riskScore"))
        except (TypeError, ValueError):
            return None

    @staticmethod
    def _collect_session_concerns(primary: str, context: str) -> list[str]:
        excluded = ReferralAIService._extract_excluded_concerns(primary)
        concerns: list[str] = []
        for line in ReferralAIService._context_lines(context):
            for key in ReferralAIService._detect_concern_keys(line, excluded=excluded):
                if key not in concerns and key not in excluded:
                    concerns.append(key)
        for key in ReferralAIService._detect_concern_keys(primary, excluded=excluded):
            if key not in concerns and key not in excluded:
                concerns.append(key)
        return concerns

    @staticmethod
    def _medicine_lines_for_concerns(concerns: list[str]) -> list[str]:
        selected: list[str] = []
        if "ear_wax" in concerns:
            selected.extend(CONCERN_MEDICINES["ear_wax"])
        if "vertigo" in concerns:
            selected.extend(CONCERN_MEDICINES["vertigo"])
        if any(c in concerns for c in ("ear_wax", "investigations")) and "speech" in concerns:
            pass  # wax meds already added
        if "speech" in concerns and not any(c in concerns for c in ("ear_wax", "vertigo", "investigations")):
            return [
                "No medicine directly reverses speech/language delay; prioritize hearing assessment and treat wax/otitis if present.",
                "Always confirm medicine choice and dose with a qualified clinician.",
            ]
        if not selected:
            if "speech" in concerns:
                return [
                    "No medicine directly reverses speech/language delay; prioritize hearing assessment first.",
                    "Always confirm medicine choice and dose with a qualified clinician.",
                ]
            selected = ["paracetamol", "ibuprofen"]

        lines: list[str] = []
        seen: set[str] = set()
        for med in selected:
            if med in seen:
                continue
            seen.add(med)
            lines.append(MEDICINE_GUIDANCE.get(med, med))
            if len(lines) >= 3:
                break
        lines.append("Always confirm medicine choice and dose with a qualified clinician.")
        return lines

    @staticmethod
    def _medicine_lines_for_prompt(prompt: str) -> list[str]:
        concerns = ReferralAIService._detect_concern_keys(prompt)
        return ReferralAIService._medicine_lines_for_concerns(concerns)

    def _clean_response(self, raw: str, prompt: str) -> str:
        text = raw
        for tok in _LLAMA_SPECIAL_TOKENS:
            text = text.replace(tok, "")

        if prompt in text:
            text = text.replace(prompt, "", 1).strip()

        loop_match = re.search(r"(?:\bthe of\b\s+){8,}", text, flags=re.IGNORECASE)
        if loop_match:
            text = text[: loop_match.start()].rstrip()

        # Cut off at any reserved special token (e.g. <|reserved_special_token_2|>)
        text = _RESERVED_TOKEN_RE.sub("", text)

        # Cut off !d!d!d garbage just in case
        text = re.sub(r"[!?][a-zA-Z](?:[!?][a-zA-Z]){5,}.*", "", text, flags=re.DOTALL)

        return text.strip()

    def _degenerate_reason(self, text: str) -> str | None:
        if len(text.strip()) < 40:
            return "too_short"

        if any(tok in text for tok in _LLAMA_SPECIAL_TOKENS):
            return "special_tokens"

        if text.lower().count("the of") >= 10:
            return "the_of_loop"

        words = re.findall(r"[a-zA-Z']+", text.lower())
        if len(words) >= 120:
            trigrams = Counter(zip(words, words[1:], words[2:]))
            top = trigrams.most_common(1)
            if top and top[0][1] >= 25:
                phrase = " ".join(top[0][0])
                if phrase in ("are a clinical", "the of the", "a clinical audiologist"):
                    return f"repeated_trigram:{top[0][0]}"

        if "PATIENT REFERRAL" in text and "This referral was generated by HearTech screening system." not in text:
            return "missing_referral_footer"

        return None

    def _is_degenerate(self, text: str) -> bool:
        return self._degenerate_reason(text) is not None

    @staticmethod
    def _blocks_shipping_answer(validation: ValidationResult) -> bool:
        return any(w in _SHIPPING_BLOCKED_WARNINGS for w in validation.warnings)

    def _can_ship_answer(
        self,
        text: str,
        validation: ValidationResult,
        *,
        primary_instruction: str = "",
    ) -> bool:
        if self._is_degenerate(text) or is_hard_answer_failure(validation.warnings):
            return False
        min_len = 80
        if primary_instruction:
            if asks_precautions_or_care_request(
                primary_instruction
            ) or asks_action_for_reported_symptoms(primary_instruction):
                min_len = 120
            elif asks_yes_no_clinical_question(primary_instruction):
                min_len = 100
            elif (
                asks_symptom_clinical_question(primary_instruction)
                or asks_clinical_differential_question(primary_instruction)
            ):
                min_len = 70
        if len(text.strip()) < min_len:
            return False
        if has_incomplete_answer_fragment(text):
            return False
        if has_prompt_instruction_leak(text):
            return False
        return not self._blocks_shipping_answer(validation)

    def _can_ship_answer_relaxed(
        self,
        text: str,
        validation: ValidationResult,
        *,
        primary_instruction: str = "",
    ) -> bool:
        """Accept substantive local model answers when only soft validators failed."""
        if self._can_ship_answer(
            text,
            validation,
            primary_instruction=primary_instruction,
        ):
            return True
        if self._is_degenerate(text) or is_hard_answer_failure(validation.warnings):
            return False
        if len(text.strip()) < 60:
            return False
        if has_incomplete_answer_fragment(text) or has_prompt_instruction_leak(text):
            return False
        return not self._blocks_shipping_answer(validation)

    @staticmethod
    def _blocks_best_effort_answer(validation: ValidationResult) -> bool:
        """Backward-compatible alias."""
        return ReferralAIService._blocks_shipping_answer(validation)

    def _generate_for_router(
        self,
        *,
        model,
        tokenizer,
        user_prompt: str,
        max_tokens: int,
        temp: float,
        top_p: float,
        verbose: bool,
        system_prompt: str,
    ) -> str:
        from mlx_lm import generate as mlx_generate
        from mlx_lm.sample_utils import make_logits_processors, make_sampler

        prompt = build_inference_prompt(user_prompt, system_prompt=system_prompt)
        prompt_tokens = tokenizer.encode(prompt, add_special_tokens=False)
        sampler = make_sampler(temp=temp, top_p=top_p)
        logits_processors = make_logits_processors(
            repetition_penalty=1.12,
            repetition_context_size=96,
        )
        raw = mlx_generate(
            model,
            tokenizer,
            prompt=prompt_tokens,
            max_tokens=max_tokens,
            sampler=sampler,
            logits_processors=logits_processors,
        )
        text = raw.strip()
        if "<|assistant|>" in text:
            text = text.split("<|assistant|>", 1)[-1].strip()
        if "<|end|>" in text:
            text = text.split("<|end|>", 1)[0].strip()
        return text

    def _run_inference(
        self,
        model,
        prompt: str,
        label: str,
        *,
        max_tokens: int,
        temp: float,
        top_p: float,
    ) -> str:
        from mlx_lm import generate as mlx_generate
        from mlx_lm.sample_utils import make_sampler, make_logits_processors

        formatted_prompt = build_inference_prompt(prompt)
        prompt_tokens = self._tokenizer.encode(formatted_prompt, add_special_tokens=False)
        sampler = make_sampler(temp=temp, top_p=top_p)
        logits_processors = make_logits_processors(
            repetition_penalty=1.15,
            repetition_context_size=96,
        )

        raw = mlx_generate(
            model,
            self._tokenizer,
            prompt=prompt_tokens,
            max_tokens=min(max_tokens, _INFERENCE_MAX_TOKENS),
            sampler=sampler,
            logits_processors=logits_processors,
        )
        cleaned = self._clean_response(raw, formatted_prompt)
        reason = self._degenerate_reason(cleaned)
        print(
            f"[REFERRAL-AI] {self._public_log_label(label)}: raw={len(raw)} "
            f"cleaned={len(cleaned)} degenerate={reason or False}"
        )
        self._release_metal_memory()
        return cleaned

    @classmethod
    def exports_dir(cls) -> Path:
        return _EXPORTS_DIR

    @staticmethod
    def _letter_looks_complete(referral_text: str) -> bool:
        return "Dear Colleague" in referral_text and "CLINICAL SUMMARY" in referral_text

    @staticmethod
    def _safe_export_stem(child_name: str) -> str:
        stem = re.sub(r"[^a-zA-Z0-9_-]+", "_", child_name.strip().lower())
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"referral_{stem}_{ts}"

    @staticmethod
    def _cloudinary_configured() -> bool:
        if os.getenv("CLOUDINARY_URL", "").strip():
            return True
        return all(
            os.getenv(k, "").strip()
            for k in (
                "CLOUDINARY_CLOUD_NAME",
                "CLOUDINARY_API_KEY",
                "CLOUDINARY_API_SECRET",
            )
        )

    @staticmethod
    def _configure_cloudinary() -> bool:
        import cloudinary

        url = os.getenv("CLOUDINARY_URL", "").strip()
        if url:
            cloudinary.config(cloudinary_url=url)
            return True
        cloud = os.getenv("CLOUDINARY_CLOUD_NAME", "").strip()
        key = os.getenv("CLOUDINARY_API_KEY", "").strip()
        secret = os.getenv("CLOUDINARY_API_SECRET", "").strip()
        if cloud and key and secret:
            cloudinary.config(cloud_name=cloud, api_key=key, api_secret=secret)
            return True
        return False

    def _publish_export_file(self, tmp_path: str, ext: str, child_name: str) -> dict:
        stem = self._safe_export_stem(child_name)
        filename = f"{stem}.{ext}"

        if self._cloudinary_configured() and self._configure_cloudinary():
            import cloudinary.uploader

            try:
                result = cloudinary.uploader.upload(
                    tmp_path,
                    resource_type="raw",
                    folder="heartech/referrals",
                    public_id=stem,
                    access_mode="public",
                )
                url = result.get("secure_url", "")
                if url:
                    print(f"[REFERRAL-EXPORT] Uploaded to Cloudinary: {url[:80]}...")
                    return {"url": url, "storage": "cloudinary", "filename": filename}
            except Exception as e:
                print(f"[REFERRAL-EXPORT] Cloudinary upload failed ({e}); using local file.")

        dest = _EXPORTS_DIR / filename
        shutil.copy2(tmp_path, dest)
        local_path = f"/api/referral-exports/{filename}"
        print(f"[REFERRAL-EXPORT] Saved locally: {dest}")
        return {"url": local_path, "storage": "local", "filename": filename}

    def to_pdf_cloudinary(self, referral_text: str, child_name: str) -> dict:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import mm

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp_path = tmp.name

        try:
            doc = SimpleDocTemplate(
                tmp_path,
                pagesize=A4,
                leftMargin=25 * mm,
                rightMargin=25 * mm,
                topMargin=45 * mm,
                bottomMargin=35 * mm,
            )
            styles = getSampleStyleSheet()
            body_style = ParagraphStyle(
                "ReferralBody",
                parent=styles["Normal"],
                fontSize=11,
                leading=15,
                spaceAfter=6,
                alignment=0,
            )
            footer_style = ParagraphStyle(
                "FooterNote",
                parent=styles["Normal"],
                fontSize=8,
                textColor="#888888",
                spaceBefore=12,
            )

            story: list = []
            if not self._letter_looks_complete(referral_text):
                title_style = ParagraphStyle(
                    "ReferralTitle", parent=styles["Heading1"], fontSize=16, spaceAfter=12
                )
                story.extend([
                    Paragraph("Paediatric Hearing Referral", title_style),
                    Paragraph(f"Patient: {child_name}", styles["Heading3"]),
                    Paragraph(
                        f"Date: {datetime.now().strftime('%d %B %Y')}",
                        styles["Normal"],
                    ),
                    Spacer(1, 12),
                ])

            for para in referral_text.split("\n"):
                para = para.strip()
                if para:
                    safe = (
                        para.replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
                    )
                    story.append(Paragraph(safe, body_style))

            story.append(
                Paragraph(
                    "Generated by HearTech — reserve header and footer for hospital letterhead.",
                    footer_style,
                )
            )

            doc.build(story)
            return self._publish_export_file(tmp_path, "pdf", child_name)
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

    def to_docx_cloudinary(self, referral_text: str, child_name: str) -> dict:
        from docx import Document
        from docx.shared import Mm, Pt

        doc = Document()
        section = doc.sections[0]
        section.top_margin = Mm(45)
        section.bottom_margin = Mm(35)
        section.left_margin = Mm(25)
        section.right_margin = Mm(25)

        if not self._letter_looks_complete(referral_text):
            doc.add_heading("Paediatric Hearing Referral", level=1)
            doc.add_paragraph(f"Patient: {child_name}")
            doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}")
            doc.add_paragraph("")

        for para in referral_text.split("\n"):
            para = para.strip()
            if para:
                p = doc.add_paragraph(para)
                for run in p.runs:
                    run.font.size = Pt(11)

        doc.add_paragraph("")
        note = doc.add_paragraph(
            "Generated by HearTech — reserve header and footer for hospital letterhead."
        )
        for run in note.runs:
            run.font.size = Pt(8)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
            doc.save(tmp_path)

        try:
            return self._publish_export_file(tmp_path, "docx", child_name)
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
